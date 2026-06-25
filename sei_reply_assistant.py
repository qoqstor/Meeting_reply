"""
Live Translation + Reply Assistant
Monitors jt-live-whisper transcript log; shows real-time ja/zh translation;
manual button triggers AI reply suggestion using configurable project profiles.

Supports: Anthropic Claude, Google Gemini, local Ollama.
Profiles: <app_dir>/profiles/*.json
Settings: <app_dir>/settings.json  (jtw_base_dir, last_provider, last_model, last_profile)

v2: Auto-trigger on sentence-end punctuation, bidirectional reply (ja+zh draft),
    multi-turn conversation history, edge-tts playback, Word/Excel transcript export.
"""
import asyncio
import json
import os
import queue
import re
import subprocess
import sys
import tempfile
import threading
import time
import tkinter as tk
from tkinter import scrolledtext, ttk, simpledialog, messagebox, filedialog
import glob
import urllib.request
from datetime import datetime, timedelta
from pathlib import Path

try:
    import pyperclip
except ImportError:
    pyperclip = None

try:
    import anthropic as _anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

try:
    from google import genai as _genai
    HAS_GEMINI = True
except ImportError:
    HAS_GEMINI = False

try:
    import edge_tts as _edge_tts
    HAS_EDGE_TTS = True
except ImportError:
    HAS_EDGE_TTS = False

try:
    from docx import Document as _DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl as _openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

# ── App directory (works for both .py script and PyInstaller .exe) ────────────
if getattr(sys, 'frozen', False):
    APP_DIR = Path(sys.executable).parent
else:
    APP_DIR = Path(__file__).parent

SETTINGS_FILE = APP_DIR / "settings.json"
PROFILES_DIR  = APP_DIR / "profiles"
PROFILES_DIR.mkdir(exist_ok=True)

JTW_EXE = APP_DIR / "jtw" / "jt-live-whisper.exe"

VERSION = "2.1.0"
GITHUB_RELEASES_API = "https://api.github.com/repos/qoqstor/Meeting_reply/releases/latest"

# ── Settings ──────────────────────────────────────────────────────────────────
def _load_settings() -> dict:
    if SETTINGS_FILE.exists():
        try:
            return json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}

def _save_settings(s: dict):
    SETTINGS_FILE.write_text(json.dumps(s, ensure_ascii=False, indent=2), encoding="utf-8")

def _check_update():
    """Background thread: compare GitHub latest release tag with VERSION.
    Pushes update_available event if newer; silently swallows all errors."""
    try:
        req = urllib.request.Request(
            GITHUB_RELEASES_API,
            headers={
                "Accept": "application/vnd.github+json",
                "User-Agent": f"SEI-Reply-Assistant/{VERSION}",
            },
        )
        with urllib.request.urlopen(req, timeout=8) as r:
            data = json.loads(r.read().decode())
        latest = data["tag_name"].lstrip("v")
        if (tuple(int(x) for x in latest.split("."))
                > tuple(int(x) for x in VERSION.split("."))):
            asset_url = next(
                (a["browser_download_url"] for a in data.get("assets", [])
                 if a["name"].lower().startswith("sei_") and a["name"].endswith(".zip")),
                None,
            )
            if asset_url:
                _ui_queue.put({
                    "type": "update_available",
                    "version": latest,
                    "url": asset_url,
                    "notes": data.get("body", "")[:150],
                })
    except Exception:
        pass

_settings = _load_settings()

_jtw_base    = Path(_settings.get("jtw_base_dir", str(APP_DIR)))
LOGS_DIR     = str(_jtw_base / "logs")
START_SCRIPT = str(_jtw_base / "start.ps1")

def _jtw_log_dirs() -> list[str]:
    dirs: list[str] = [LOGS_DIR]
    if JTW_EXE.exists():
        jtw_dir = JTW_EXE.parent
        dirs.append(str(jtw_dir / "logs"))
        dirs.append(str(jtw_dir / "_internal" / "logs"))
    return list(dict.fromkeys(dirs))

# ── Constants ─────────────────────────────────────────────────────────────────
MAX_SENTENCES         = 30
MAX_MINUTES           = 10
MAX_HISTORY_TURNS     = 10   # keep last N user/assistant pairs
OLLAMA_HOST           = "http://localhost:11434"
AUTO_TRIGGER_DEBOUNCE = 8.0  # seconds after last sentence-ending punctuation

JTW_MODES = [
    ("ja_zh",  "日中雙向"),
    ("ja2zh",  "日翻中"),
    ("zh2ja",  "中翻日"),
    ("en_zh",  "英中雙向"),
    ("en2zh",  "英翻中"),
    ("zh2en",  "中翻英"),
]

ANTHROPIC_MODELS = ["claude-haiku-4-5-20251001", "claude-sonnet-4-6"]
GEMINI_MODELS    = ["gemini-2.0-flash-lite", "gemini-2.5-flash-preview-05-20",
                    "gemini-2.0-flash", "gemini-2.5-pro-preview-06-05"]
OLLAMA_MODELS    = ["qwen2.5:7b", "qwen2.5:14b", "llama3.1:8b"]

TTS_VOICES = {
    "ja": "ja-JP-NanamiNeural",
    "zh": "zh-TW-HsiaoChenNeural",
    "en": "en-US-JennyNeural",
}

SENTENCE_END_RE = re.compile(r'[。！？!?]')

DEFAULT_PROFILE = {
    "name": "Default",
    "description": "通用即時回覆助理",
    "reply_lang": "ja",
    "system_prompt": (
        "あなたは会議の回答アシスタントです。相手の発言に対して丁寧な敬語で3〜4文の返答を提案してください。\n"
        "以下のJSON形式のみで出力してください（余分なテキスト不要）：\n"
        '{"ja_reply": "日本語の返答（3〜4文）", "zh_reply": "中文草稿（對應翻譯）"}'
    )
}

# ── Profile management ────────────────────────────────────────────────────────
def load_profiles() -> dict[str, dict]:
    profiles = {}
    for f in sorted(PROFILES_DIR.glob("*.json")):
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            profiles[data["name"]] = data
        except Exception:
            pass
    if not profiles:
        profiles["Default"] = DEFAULT_PROFILE
    return profiles

def save_profile(profile: dict):
    name = profile["name"]
    safe = re.sub(r'[^\w\-]', '_', name)
    path = PROFILES_DIR / f"{safe}.json"
    path.write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")

# ── Shared state ──────────────────────────────────────────────────────────────
_ui_queue: queue.Queue = queue.Queue()
_selected_model    = None  # tk.StringVar
_selected_provider = None  # tk.StringVar
_current_profile   = None  # dict

_recent_lines: list[tuple[str, str, str]] = []
_lines_lock      = threading.Lock()
_monitor_started = False

_conversation_history: list[dict] = []
_history_lock = threading.Lock()
_tts_lock     = threading.Lock()

# ── AI helpers ────────────────────────────────────────────────────────────────
def _build_user_message(ja_block: str, zh_block: str) -> str:
    return (
        f"以下は相手の発言です（最大{MAX_SENTENCES}文/{MAX_MINUTES}分以内）：\n\n"
        f"{ja_block}\n\n参考訳：\n{zh_block}\n\n"
        "日本語の返答と中文草稿を以下のJSON形式で出力してください（余分なテキスト不要）：\n"
        '{"ja_reply": "日本語の返答（3〜4文）", "zh_reply": "中文草稿（對應翻譯）"}'
    )

def _parse_bilingual(text: str) -> tuple[str, str]:
    """Extract ja_reply / zh_reply from JSON response; fallback to raw text."""
    try:
        m = re.search(r'\{[^{}]*"ja_reply"[^{}]*\}', text, re.DOTALL)
        if m:
            data = json.loads(m.group())
            return data.get("ja_reply", "").strip(), data.get("zh_reply", "").strip()
    except (json.JSONDecodeError, AttributeError):
        pass
    return text.strip(), ""

def call_ollama_stream(model: str, messages: list[dict]):
    payload = json.dumps({
        "model": model,
        "messages": messages,
        "stream": True,
        "options": {"num_predict": 500, "temperature": 0.7}
    }).encode("utf-8")
    req = urllib.request.Request(
        f"{OLLAMA_HOST}/api/chat", data=payload,
        headers={"Content-Type": "application/json"}, method="POST"
    )
    buf = ""
    last_yield = time.monotonic()
    with urllib.request.urlopen(req, timeout=300) as resp:
        for raw_line in resp:
            line = raw_line.decode("utf-8").strip()
            if not line:
                continue
            try:
                data = json.loads(line)
                chunk = data.get("message", {}).get("content", "")
                if chunk:
                    buf += chunk
                    if time.monotonic() - last_yield >= 0.3:
                        yield buf
                        last_yield = time.monotonic()
                if data.get("done"):
                    if buf:
                        yield buf
                    break
            except json.JSONDecodeError:
                continue

def call_ai(api_key: str, provider: str, model: str,
            ja_block: str, zh_block: str, system: str,
            history: list[dict]) -> tuple[str, str]:
    """Call AI; stream chunks via _ui_queue; return (ja_reply, zh_reply)."""
    user_msg  = _build_user_message(ja_block, zh_block)
    hist_msgs = [m for m in history if m["role"] in ("user", "assistant")]
    full_text = ""

    if provider == "Ollama":
        messages = ([{"role": "system", "content": system}]
                    + hist_msgs
                    + [{"role": "user", "content": user_msg}])
        for chunk in call_ollama_stream(model, messages):
            full_text = chunk
            _ui_queue.put({"type": "reply_chunk", "text": chunk})

    elif provider == "Anthropic":
        messages = hist_msgs + [{"role": "user", "content": user_msg}]
        client = _anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(
            model=model, max_tokens=600, system=system,
            messages=messages)
        full_text = msg.content[0].text.strip()

    else:  # Gemini
        hist_text = "".join(
            f"{'User' if m['role']=='user' else 'Assistant'}: {m['content']}\n\n"
            for m in hist_msgs
        )
        content = hist_text + f"User: {user_msg}"
        client = _genai.Client(api_key=api_key)
        resp = client.models.generate_content(
            model=model, contents=content,
            config={"system_instruction": system, "max_output_tokens": 600})
        full_text = resp.text.strip()

    return _parse_bilingual(full_text)

# ── TTS ───────────────────────────────────────────────────────────────────────
async def _tts_save_async(text: str, voice: str) -> str:
    tmp = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
    tmp_path = tmp.name
    tmp.close()
    communicate = _edge_tts.Communicate(text, voice)
    await communicate.save(tmp_path)
    return tmp_path

def play_tts(text: str, lang: str = "ja"):
    if not HAS_EDGE_TTS:
        _ui_queue.put({"type": "status", "text": "TTS 未安裝：pip install edge-tts"})
        return
    if not text.strip():
        return
    voice = TTS_VOICES.get(lang, TTS_VOICES["ja"])

    def _run():
        if not _tts_lock.acquire(blocking=False):
            return
        try:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            tmp_path = loop.run_until_complete(_tts_save_async(text, voice))
            loop.close()
            os.startfile(tmp_path)
            time.sleep(90)
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
        except Exception as e:
            _ui_queue.put({"type": "status", "text": f"TTS 錯誤：{e}"})
        finally:
            _tts_lock.release()

    threading.Thread(target=_run, daemon=True).start()

# ── Export ────────────────────────────────────────────────────────────────────
def export_excel(path: str, records: list[dict]):
    wb = _openpyxl.Workbook()
    ws = wb.active
    ws.title = "翻譯記錄"
    from openpyxl.styles import Font, PatternFill
    headers = ["時間", "日文（原文）", "中文（翻譯）", "AI 日文回覆", "中文草稿"]
    ws.append(headers)
    blue_fill = PatternFill("solid", fgColor="2563EB")
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = blue_fill
    for r in records:
        ws.append([r.get("ts",""), r.get("ja",""), r.get("zh",""),
                   r.get("ja_reply",""), r.get("zh_reply","")])
    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=0)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 60)
    wb.save(path)

def export_word(path: str, records: list[dict]):
    doc = _DocxDocument()
    doc.add_heading("即時翻譯 + 回覆記錄", 0)
    doc.add_paragraph(f"匯出時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"共 {len(records)} 筆記錄")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["時間", "日文", "中文", "AI回覆(日)", "草稿(中)"]):
        hdr[i].text = h
    for r in records:
        row = table.add_row().cells
        row[0].text = r.get("ts", "")
        row[1].text = r.get("ja", "")
        row[2].text = r.get("zh", "")
        row[3].text = r.get("ja_reply", "")
        row[4].text = r.get("zh_reply", "")
    doc.save(path)

# ── Log monitor ───────────────────────────────────────────────────────────────
LINE_RE = re.compile(r'^\[(\d{2}:\d{2}:\d{2})\]\s*(?:[▶◀]\s*)?\[([日中])\]\s*(.+)$')

def _find_latest_log() -> str | None:
    files = []
    for d in _jtw_log_dirs():
        files.extend(glob.glob(os.path.join(d, "*逐字稿_*.txt")))
    return max(files, key=os.path.getmtime) if files else None

def log_monitor():
    current_log = None
    file_pos    = 0
    pending_zh: dict[str, str] = {}
    _ui_queue.put({"type": "status", "text": "尋找 log 檔…"})

    while True:
        latest = _find_latest_log()
        if latest != current_log:
            current_log = latest
            file_pos = 0
            if current_log:
                file_pos = os.path.getsize(current_log)
                _ui_queue.put({"type": "status",
                               "text": f"監控：{os.path.basename(current_log)}"})
            else:
                _ui_queue.put({"type": "status", "text": "等待 jt-live-whisper 啟動…"})
                time.sleep(3)
                continue

        try:
            sz = os.path.getsize(current_log)
        except OSError:
            time.sleep(1)
            continue

        if sz > file_pos:
            try:
                with open(current_log, "r", encoding="utf-8", errors="replace") as f:
                    f.seek(file_pos)
                    new_content = f.read()
                file_pos = sz
            except OSError:
                time.sleep(1)
                continue

            for line in new_content.splitlines():
                m = LINE_RE.match(line.strip())
                if not m:
                    continue
                ts, lang, text = m.group(1), m.group(2), m.group(3).strip()

                if lang == "日":
                    zh = pending_zh.pop(ts, "")
                    with _lines_lock:
                        _recent_lines.append((ts, text, zh))
                        if len(_recent_lines) > 100:
                            _recent_lines.pop(0)
                    _ui_queue.put({
                        "type": "transcript_add",
                        "ts": ts, "ja": text, "zh": zh,
                        "has_sentence_end": bool(SENTENCE_END_RE.search(text))
                    })

                elif lang == "中":
                    with _lines_lock:
                        for i in range(len(_recent_lines) - 1, -1, -1):
                            if _recent_lines[i][0] == ts:
                                t_, ja_, _ = _recent_lines[i]
                                _recent_lines[i] = (t_, ja_, text)
                                break
                        else:
                            pending_zh[ts] = text
                    _ui_queue.put({"type": "transcript_update_zh",
                                   "ts": ts, "zh": text})

        time.sleep(0.4)

def _start_monitor():
    global _monitor_started
    if _monitor_started:
        return
    _monitor_started = True
    threading.Thread(target=log_monitor, daemon=True).start()

def _get_recent_context() -> tuple[str, str]:
    cutoff = (datetime.now() - timedelta(minutes=MAX_MINUTES)).strftime("%H:%M:%S")
    with _lines_lock:
        cands = [(ts, ja, zh) for ts, ja, zh in _recent_lines
                 if ts >= cutoff and ja]
    sel = cands[-MAX_SENTENCES:]
    return (
        "\n".join(f"[{ts}] {ja}" for ts, ja, _ in sel),
        "\n".join(f"[{ts}] {zh}" for ts, _, zh in sel),
    )

# ── UI ────────────────────────────────────────────────────────────────────────
class App:
    BG     = "#1e1e2e"
    CARD   = "#2a2a3e"
    ACCENT = "#7c3aed"
    FG     = "#cdd6f4"
    MUTED  = "#6c7086"
    GREEN  = "#a6e3a1"
    BLUE   = "#89dceb"
    ORANGE = "#fab387"
    PINK   = "#f5c2e7"

    def __init__(self, root: tk.Tk):
        global _selected_model, _selected_provider, _current_profile
        self.root = root
        self.root.title("即時翻譯 + 回覆助理 v2")
        self.root.geometry("960x940")
        self.root.configure(bg=self.BG)
        self.root.attributes("-topmost", True)

        self._api_key    = ""
        self._generating = False
        self._profiles   = load_profiles()
        _current_profile = list(self._profiles.values())[0]

        last_provider = _settings.get("last_provider", "Ollama")
        last_model    = _settings.get("last_model", OLLAMA_MODELS[0])
        last_profile  = _settings.get("last_profile", _current_profile["name"])

        _selected_provider = tk.StringVar(value=last_provider)
        _selected_model    = tk.StringVar(value=last_model)
        self._profile_var  = tk.StringVar(value=last_profile)

        if last_profile in self._profiles:
            _current_profile = self._profiles[last_profile]

        self._ts_linemap: dict[str, int] = {}
        self._trans_line = 1

        # v2 state
        self._auto_reply_var  = tk.BooleanVar(value=False)
        self._tts_auto_var    = tk.BooleanVar(value=False)
        self._tts_lang_var    = tk.StringVar(value="ja")
        self._last_sentence_end: float = 0.0
        self._reply_sessions: list[dict] = []
        self._last_ja_reply = ""
        self._last_zh_reply = ""

        self._build()
        self._poll()
        self._check_auto_trigger()
        self.root.after(300, self._check_jtw_dir)

        # Check for newer release in background (silent if no internet)
        threading.Thread(target=_check_update, daemon=True).start()

    # ── jt-live-whisper availability check ───────────────────────────────────
    def _check_jtw_dir(self):
        global _settings, _jtw_base, LOGS_DIR, START_SCRIPT

        if JTW_EXE.exists():
            self._status_var.set(f"jt-live-whisper (bundled): {JTW_EXE.parent}")
            self._jtw_dir_var.set(str(JTW_EXE.parent))
            self.root.after(100, _start_monitor)
            return

        if Path(START_SCRIPT).exists():
            self._status_var.set(f"jt-live-whisper: {_jtw_base}")
            self.root.after(100, _start_monitor)
            return

        messagebox.showinfo(
            "設定 jt-live-whisper 目錄",
            "首次使用請選擇 jt-live-whisper 安裝目錄\n"
            "（包含 start.ps1 的資料夾，例如 C:\\jt-live-whisper）",
            parent=self.root
        )
        chosen = filedialog.askdirectory(
            title="選擇 jt-live-whisper 安裝目錄",
            parent=self.root
        )
        if not chosen:
            self._status_var.set("⚠ 未選擇目錄，部分功能不可用")
            self.root.after(100, _start_monitor)
            return

        _jtw_base    = Path(chosen)
        LOGS_DIR     = str(_jtw_base / "logs")
        START_SCRIPT = str(_jtw_base / "start.ps1")
        _settings["jtw_base_dir"] = str(_jtw_base)
        _save_settings(_settings)
        self._status_var.set(f"jt-live-whisper: {_jtw_base}")
        self.root.after(100, _start_monitor)

    # ── Build UI ──────────────────────────────────────────────────────────────
    def _build(self):
        # Title bar
        hdr = tk.Frame(self.root, bg=self.ACCENT, pady=5)
        hdr.pack(fill="x")
        tk.Label(hdr, text="即時翻譯 + 回覆助理 v2",
                 bg=self.ACCENT, fg="white",
                 font=("Arial", 12, "bold")).pack(side="left", padx=10)
        self._jtw_dir_var = tk.StringVar(value=str(_jtw_base))
        tk.Label(hdr, textvariable=self._jtw_dir_var,
                 bg=self.ACCENT, fg="#ddd6fe", font=("Arial", 8)).pack(side="right", padx=6)
        tk.Button(hdr, text="📁", command=self._change_jtw_dir,
                  bg=self.ACCENT, fg="white", relief="flat",
                  font=("Arial", 10), cursor="hand2").pack(side="right")

        # Row 0: jt-live-whisper launcher
        launch_f = tk.Frame(self.root, bg="#0f0f1a", pady=5, padx=10)
        launch_f.pack(fill="x", padx=6, pady=(6, 0))
        tk.Label(launch_f, text="翻譯模式:", bg="#0f0f1a",
                 fg=self.MUTED, font=("Arial", 9)).pack(side="left")
        self._jtw_mode_var = tk.StringVar(value="ja_zh")
        ttk.Combobox(launch_f, textvariable=self._jtw_mode_var,
                     values=[f"{v}  {l}" for v, l in JTW_MODES],
                     width=14, state="readonly").pack(side="left", padx=(4, 8))
        self._launch_btn = tk.Button(
            launch_f, text="🎙 開啟即時翻譯", command=self._launch_jtw,
            bg="#0369a1", fg="white", font=("Arial", 10, "bold"),
            relief="flat", padx=12, pady=3, cursor="hand2")
        self._launch_btn.pack(side="left")
        tk.Button(launch_f, text="⏹ 停止翻譯", command=self._stop_jtw,
                  bg="#7f1d1d", fg="white", font=("Arial", 9), relief="flat",
                  padx=10, pady=3, cursor="hand2").pack(side="left", padx=6)
        self._jtw_status_var = tk.StringVar(value="")
        self._jtw_status_lbl = tk.Label(
            launch_f, textvariable=self._jtw_status_var,
            bg="#0f0f1a", fg=self.MUTED, font=("Arial", 9))
        self._jtw_status_lbl.pack(side="left", padx=8)
        self._jtw_proc = None
        self._check_jtw_status()

        # Row 1: Profile
        r0 = tk.Frame(self.root, bg=self.CARD, pady=5, padx=10)
        r0.pack(fill="x", padx=6, pady=(6, 0))
        tk.Label(r0, text="專案情境:", bg=self.CARD, fg=self.FG,
                 font=("Arial", 9)).pack(side="left")
        self._profile_cb = ttk.Combobox(
            r0, textvariable=self._profile_var,
            values=list(self._profiles.keys()), width=20, state="readonly")
        self._profile_cb.pack(side="left", padx=6)
        self._profile_cb.bind("<<ComboboxSelected>>", self._on_profile_change)
        tk.Button(r0, text="編輯情境", command=self._edit_profile,
                  bg=self.CARD, fg=self.ORANGE, relief="flat",
                  padx=8, cursor="hand2").pack(side="left", padx=4)
        tk.Button(r0, text="新增情境", command=self._new_profile,
                  bg=self.CARD, fg=self.GREEN, relief="flat",
                  padx=8, cursor="hand2").pack(side="left", padx=4)
        self._desc_var = tk.StringVar(value=_current_profile.get("description", ""))
        tk.Label(r0, textvariable=self._desc_var, bg=self.CARD,
                 fg=self.MUTED, font=("Arial", 8)).pack(side="left", padx=8)

        # Row 2: Provider + API Key
        r1 = tk.Frame(self.root, bg=self.CARD, pady=5, padx=10)
        r1.pack(fill="x", padx=6, pady=(2, 0))
        tk.Label(r1, text="Provider:", bg=self.CARD, fg=self.FG,
                 font=("Arial", 9)).pack(side="left")
        pcb = ttk.Combobox(r1, textvariable=_selected_provider,
                            values=["Ollama", "Anthropic", "Gemini"],
                            width=10, state="readonly")
        pcb.pack(side="left", padx=(4, 12))
        pcb.bind("<<ComboboxSelected>>", self._on_provider_change)
        tk.Label(r1, text="API Key:", bg=self.CARD, fg=self.FG,
                 font=("Arial", 9)).pack(side="left")
        self._key_var = tk.StringVar()
        tk.Entry(r1, textvariable=self._key_var, show="*",
                 width=38, bg="#313244", fg=self.FG,
                 insertbackground=self.FG, relief="flat").pack(side="left", padx=6)
        tk.Button(r1, text="連線", command=self._connect,
                  bg=self.ACCENT, fg="white", relief="flat",
                  padx=10, cursor="hand2").pack(side="left")

        # Row 3: Model + history counter
        r2 = tk.Frame(self.root, bg=self.CARD, pady=3, padx=10)
        r2.pack(fill="x", padx=6)
        tk.Label(r2, text="Model:", bg=self.CARD, fg=self.MUTED,
                 font=("Arial", 9)).pack(side="left")
        p = _selected_provider.get()
        initial_models = (ANTHROPIC_MODELS if p == "Anthropic"
                          else GEMINI_MODELS if p == "Gemini"
                          else OLLAMA_MODELS)
        self._model_cb = ttk.Combobox(
            r2, textvariable=_selected_model,
            values=initial_models, width=28, state="readonly")
        self._model_cb.pack(side="left", padx=6)
        self._history_count_var = tk.StringVar(value="歷史：0 輪")
        tk.Label(r2, textvariable=self._history_count_var,
                 bg=self.CARD, fg=self.MUTED, font=("Arial", 8)).pack(side="left", padx=8)
        tk.Button(r2, text="清除對話歷史", command=self._clear_history,
                  bg=self.CARD, fg=self.ORANGE, relief="flat",
                  padx=8, cursor="hand2").pack(side="left")

        # Status bar
        self._status_var = tk.StringVar(value="啟動中…")
        tk.Label(self.root, textvariable=self._status_var,
                 bg=self.BG, fg=self.MUTED, font=("Arial", 9),
                 anchor="w", padx=10).pack(fill="x", pady=(3, 0))

        # Main resizable area: transcript (top) + reply (bottom) split by sash
        main_pane = ttk.PanedWindow(self.root, orient="vertical")
        main_pane.pack(fill="both", expand=True, padx=6, pady=(5, 0))

        # ── Top pane: live transcript (ja left, zh right) ──────────────────────
        trans_outer = tk.LabelFrame(
            main_pane, text="  即時翻譯  ",
            bg=self.BG, fg=self.GREEN, font=("Arial", 8, "bold"),
            padx=4, pady=4)
        main_pane.add(trans_outer, weight=3)
        trans_inner = tk.Frame(trans_outer, bg=self.BG)
        trans_inner.pack(fill="both", expand=True)

        ja_f = tk.Frame(trans_inner, bg=self.BG)
        ja_f.pack(side="left", fill="both", expand=True, padx=(0, 2))
        tk.Label(ja_f, text="🗣 日文（原文）", bg=self.BG, fg=self.GREEN,
                 font=("Arial", 9, "bold")).pack(anchor="w")
        self._trans_ja = scrolledtext.ScrolledText(
            ja_f, bg=self.CARD, fg=self.GREEN,
            font=("Meiryo", 10), relief="flat", wrap="word", state="disabled")
        self._trans_ja.pack(fill="both", expand=True)

        zh_f = tk.Frame(trans_inner, bg=self.BG)
        zh_f.pack(side="left", fill="both", expand=True, padx=(2, 0))
        tk.Label(zh_f, text="💬 中文（翻譯）", bg=self.BG, fg=self.FG,
                 font=("Arial", 9, "bold")).pack(anchor="w")
        self._trans_zh = scrolledtext.ScrolledText(
            zh_f, bg=self.CARD, fg=self.FG,
            font=("Arial", 10), relief="flat", wrap="word", state="disabled")
        self._trans_zh.pack(fill="both", expand=True)

        for w in (self._trans_ja, self._trans_zh):
            w.tag_configure("ts",      foreground=self.MUTED, font=("Arial", 8))
            w.tag_configure("text_ja", foreground=self.GREEN, font=("Meiryo", 10))
            w.tag_configure("text_zh", foreground=self.FG,    font=("Arial", 10))

        # ── Bottom pane: generate button + reply areas ─────────────────────────
        bottom_frame = tk.Frame(main_pane, bg=self.BG)
        main_pane.add(bottom_frame, weight=2)

        # Generate button + auto-trigger
        gen_f = tk.Frame(bottom_frame, bg=self.BG, pady=4)
        gen_f.pack(fill="x")
        self._gen_btn = tk.Button(
            gen_f,
            text=f"▶  生成回覆建議（最近 {MAX_SENTENCES} 句 / {MAX_MINUTES} 分鐘）",
            command=self._generate,
            bg="#059669", fg="white",
            font=("Arial", 11, "bold"),
            relief="flat", pady=6, cursor="hand2")
        self._gen_btn.pack(side="left", fill="x", expand=True)
        auto_f = tk.Frame(gen_f, bg=self.BG, padx=8)
        auto_f.pack(side="right")
        tk.Checkbutton(
            auto_f, text="自動觸發", variable=self._auto_reply_var,
            bg=self.BG, fg=self.GREEN, selectcolor=self.CARD,
            activebackground=self.BG, font=("Arial", 9)).pack(anchor="w")
        tk.Label(auto_f, text="句尾標點後 8 秒",
                 bg=self.BG, fg=self.MUTED, font=("Arial", 7)).pack(anchor="w")

        # Reply areas: ja (left) + zh draft (right)
        reply_outer = tk.LabelFrame(
            bottom_frame, text="  建議回覆  ",
            bg=self.BG, fg=self.BLUE, font=("Arial", 8, "bold"),
            padx=4, pady=4)
        reply_outer.pack(fill="both", expand=True, pady=(2, 2))
        reply_inner = tk.Frame(reply_outer, bg=self.BG)
        reply_inner.pack(fill="both", expand=True)

        ja_reply_f = tk.Frame(reply_inner, bg=self.BG)
        ja_reply_f.pack(side="left", fill="both", expand=True, padx=(0, 2))
        ja_hdr = tk.Frame(ja_reply_f, bg=self.BG)
        ja_hdr.pack(fill="x")
        tk.Label(ja_hdr, text="🤖 AI 日文回覆", bg=self.BG, fg=self.BLUE,
                 font=("Arial", 9, "bold")).pack(side="left")
        tk.Button(ja_hdr, text="▶ TTS",
                  command=lambda: self._play_reply_tts("ja"),
                  bg="#1e3a5f", fg=self.BLUE, relief="flat",
                  padx=6, font=("Arial", 8), cursor="hand2").pack(side="right")
        tk.Button(ja_hdr, text="複製",
                  command=lambda: self._copy_reply("ja"),
                  bg=self.CARD, fg=self.BLUE, relief="flat",
                  padx=6, font=("Arial", 8), cursor="hand2").pack(side="right", padx=2)
        self._reply_ja = scrolledtext.ScrolledText(
            ja_reply_f, bg=self.CARD, fg=self.BLUE,
            font=("Meiryo", 11), relief="flat", wrap="word", state="disabled")
        self._reply_ja.pack(fill="both", expand=True)

        zh_reply_f = tk.Frame(reply_inner, bg=self.BG)
        zh_reply_f.pack(side="left", fill="both", expand=True, padx=(2, 0))
        zh_hdr = tk.Frame(zh_reply_f, bg=self.BG)
        zh_hdr.pack(fill="x")
        tk.Label(zh_hdr, text="📝 中文草稿（您可說）", bg=self.BG, fg=self.PINK,
                 font=("Arial", 9, "bold")).pack(side="left")
        tk.Button(zh_hdr, text="▶ TTS",
                  command=lambda: self._play_reply_tts("zh"),
                  bg="#3b1a3e", fg=self.PINK, relief="flat",
                  padx=6, font=("Arial", 8), cursor="hand2").pack(side="right")
        tk.Button(zh_hdr, text="複製",
                  command=lambda: self._copy_reply("zh"),
                  bg=self.CARD, fg=self.PINK, relief="flat",
                  padx=6, font=("Arial", 8), cursor="hand2").pack(side="right", padx=2)
        self._reply_zh = scrolledtext.ScrolledText(
            zh_reply_f, bg=self.CARD, fg=self.PINK,
            font=("Arial", 11), relief="flat", wrap="word", state="disabled")
        self._reply_zh.pack(fill="both", expand=True)

        # Bottom controls
        btn_row = tk.Frame(bottom_frame, bg=self.BG, pady=3)
        btn_row.pack(fill="x")
        tk.Checkbutton(
            btn_row, text="自動播放 TTS", variable=self._tts_auto_var,
            bg=self.BG, fg=self.MUTED, selectcolor=self.CARD,
            activebackground=self.BG).pack(side="left")
        tk.Label(btn_row, text="語言:", bg=self.BG, fg=self.MUTED,
                 font=("Arial", 8)).pack(side="left", padx=(6, 2))
        ttk.Combobox(btn_row, textvariable=self._tts_lang_var,
                     values=["ja", "zh"], width=4, state="readonly").pack(side="left")
        if not HAS_EDGE_TTS:
            tk.Label(btn_row, text="(pip install edge-tts)",
                     bg=self.BG, fg=self.ORANGE, font=("Arial", 7)).pack(side="left", padx=4)
        tk.Button(btn_row, text="匯出逐字稿", command=self._export_transcript,
                  bg=self.CARD, fg=self.GREEN, relief="flat",
                  padx=10, cursor="hand2").pack(side="right")
        tk.Button(btn_row, text="清除記錄", command=self._clear,
                  bg=self.CARD, fg=self.MUTED, relief="flat",
                  padx=10, cursor="hand2").pack(side="right", padx=6)
        self._ontop_var = tk.BooleanVar(value=True)
        tk.Checkbutton(
            btn_row, text="視窗置頂", variable=self._ontop_var,
            command=lambda: self.root.attributes("-topmost", self._ontop_var.get()),
            bg=self.BG, fg=self.MUTED, selectcolor=self.CARD,
            activebackground=self.BG).pack(side="right", padx=6)

    # ── Change jt-live-whisper directory ──────────────────────────────────────
    def _change_jtw_dir(self):
        global _settings, _jtw_base, LOGS_DIR, START_SCRIPT
        chosen = filedialog.askdirectory(
            title="選擇 jt-live-whisper 安裝目錄",
            parent=self.root, initialdir=str(_jtw_base))
        if not chosen:
            return
        _jtw_base    = Path(chosen)
        LOGS_DIR     = str(_jtw_base / "logs")
        START_SCRIPT = str(_jtw_base / "start.ps1")
        _settings["jtw_base_dir"] = str(_jtw_base)
        _save_settings(_settings)
        self._jtw_dir_var.set(str(_jtw_base))
        self._status_var.set(f"目錄已更新：{_jtw_base}")

    # ── jt-live-whisper launcher ──────────────────────────────────────────────
    def _launch_jtw(self):
        mode = self._jtw_mode_var.get().split()[0]
        if self._jtw_proc and self._jtw_proc.poll() is None:
            self._jtw_status_var.set("⚠ 已在執行中")
            return
        try:
            if JTW_EXE.exists():
                self._jtw_proc = subprocess.Popen(
                    [str(JTW_EXE), "--_tm_mode", "--mode", mode],
                    cwd=str(JTW_EXE.parent),
                    creationflags=subprocess.CREATE_NEW_CONSOLE)
            else:
                self._jtw_proc = subprocess.Popen(
                    ["powershell.exe", "-NoProfile", "-ExecutionPolicy", "Bypass",
                     "-File", START_SCRIPT, "--mode", mode],
                    cwd=str(_jtw_base),
                    creationflags=subprocess.CREATE_NEW_CONSOLE)
            self._jtw_status_var.set(f"▶ 執行中 (pid {self._jtw_proc.pid}) [{mode}]")
            self._jtw_status_lbl.config(fg=self.GREEN)
        except Exception as e:
            self._jtw_status_var.set(f"啟動失敗：{e}")
            self._jtw_status_lbl.config(fg=self.ORANGE)

    def _stop_jtw(self):
        if self._jtw_proc and self._jtw_proc.poll() is None:
            self._jtw_proc.terminate()
            self._jtw_status_var.set("■ 已停止")
            self._jtw_status_lbl.config(fg=self.MUTED)
        else:
            kill_cmds = [
                "Get-Process 'jt-live-whisper' -ErrorAction SilentlyContinue | Stop-Process -Force",
                "Get-Process python -ErrorAction SilentlyContinue |"
                " Where-Object {$_.CommandLine -like '*translate_meeting*'} |"
                " Stop-Process -Force",
            ]
            for cmd in kill_cmds:
                subprocess.run(["powershell.exe", "-NoProfile", "-Command", cmd],
                               capture_output=True)
            self._jtw_status_var.set("■ 已停止")
            self._jtw_status_lbl.config(fg=self.MUTED)

    def _check_jtw_status(self):
        running = False
        if self._jtw_proc and self._jtw_proc.poll() is None:
            running = True
        else:
            try:
                ps_cmd = (
                    "$n = 0;"
                    "$n += (Get-Process 'jt-live-whisper' -ErrorAction SilentlyContinue | Measure-Object).Count;"
                    "$n += (Get-Process python -ErrorAction SilentlyContinue |"
                    " Where-Object {$_.CommandLine -like '*translate_meeting*'} | Measure-Object).Count;"
                    "Write-Output $n"
                )
                r = subprocess.run(
                    ["powershell.exe", "-NoProfile", "-Command", ps_cmd],
                    capture_output=True, text=True, timeout=2)
                if r.stdout.strip() and int(r.stdout.strip()) > 0:
                    running = True
            except Exception:
                pass
        if running:
            if not self._jtw_status_var.get().startswith("▶"):
                self._jtw_status_var.set("▶ 翻譯執行中")
                self._jtw_status_lbl.config(fg=self.GREEN)
        else:
            if self._jtw_status_var.get().startswith("▶"):
                self._jtw_status_var.set("● 未執行")
                self._jtw_status_lbl.config(fg=self.MUTED)
            elif not self._jtw_status_var.get():
                self._jtw_status_var.set("● 未執行")
        self.root.after(3000, self._check_jtw_status)

    def _on_profile_change(self, _=None):
        global _current_profile
        name = self._profile_var.get()
        _current_profile = self._profiles.get(name, DEFAULT_PROFILE)
        self._desc_var.set(_current_profile.get("description", ""))
        _settings["last_profile"] = name
        _save_settings(_settings)

    def _on_provider_change(self, _=None):
        p = _selected_provider.get()
        models = (ANTHROPIC_MODELS if p == "Anthropic"
                  else GEMINI_MODELS if p == "Gemini"
                  else OLLAMA_MODELS)
        self._model_cb.config(values=models)
        _selected_model.set(models[0])
        _settings["last_provider"] = p
        _settings["last_model"]    = models[0]
        _save_settings(_settings)

    def _connect(self):
        key = self._key_var.get().strip()
        p   = _selected_provider.get()
        if p != "Ollama" and not key:
            self._status_var.set("請輸入 API Key（Ollama 免填）")
            return
        self._api_key = key
        self._status_var.set(f"啟動 log 監控… [{p}]")
        _settings["last_model"] = _selected_model.get()
        _save_settings(_settings)
        _start_monitor()

    def _edit_profile(self):
        global _current_profile
        dlg = tk.Toplevel(self.root)
        dlg.title(f"編輯情境：{_current_profile['name']}")
        dlg.geometry("700x420")
        dlg.configure(bg=self.BG)
        dlg.grab_set()
        tk.Label(dlg, text="描述：", bg=self.BG, fg=self.FG,
                 font=("Arial", 9)).pack(anchor="w", padx=10, pady=(10, 0))
        desc_var = tk.StringVar(value=_current_profile.get("description", ""))
        tk.Entry(dlg, textvariable=desc_var, bg=self.CARD, fg=self.FG,
                 insertbackground=self.FG, relief="flat", width=80).pack(padx=10, fill="x")
        tk.Label(dlg, text="System Prompt（給 AI 的專案背景）：",
                 bg=self.BG, fg=self.FG, font=("Arial", 9)).pack(
                     anchor="w", padx=10, pady=(8, 0))
        txt = scrolledtext.ScrolledText(
            dlg, height=12, bg=self.CARD, fg=self.FG,
            font=("Consolas", 10), relief="flat", wrap="word")
        txt.pack(padx=10, fill="both", expand=True)
        txt.insert("1.0", _current_profile.get("system_prompt", ""))

        def _save():
            global _current_profile
            _current_profile["description"]   = desc_var.get().strip()
            _current_profile["system_prompt"] = txt.get("1.0", "end").strip()
            self._profiles[_current_profile["name"]] = _current_profile
            save_profile(_current_profile)
            self._desc_var.set(_current_profile["description"])
            dlg.destroy()

        tk.Button(dlg, text="儲存", command=_save,
                  bg=self.ACCENT, fg="white", relief="flat",
                  padx=14, pady=4, cursor="hand2").pack(pady=8)

    def _new_profile(self):
        name = simpledialog.askstring("新增情境", "輸入情境名稱（英數）：",
                                      parent=self.root)
        if not name:
            return
        name = name.strip()
        profile = {
            "name": name,
            "description": name,
            "reply_lang": "ja",
            "system_prompt": (
                "あなたは会議の回答アシスタントです。丁寧な敬語で3〜4文の返答を提案してください。\n"
                "以下のJSON形式のみで出力してください：\n"
                '{"ja_reply": "日本語の返答...", "zh_reply": "中文草稿..."}'
            )
        }
        self._profiles[name] = profile
        save_profile(profile)
        self._profile_cb.config(values=list(self._profiles.keys()))
        self._profile_var.set(name)
        self._on_profile_change()
        self._edit_profile()

    # ── Auto-trigger (poll every 500ms) ──────────────────────────────────────
    def _check_auto_trigger(self):
        if (self._auto_reply_var.get()
                and self._last_sentence_end > 0
                and not self._generating
                and time.monotonic() - self._last_sentence_end >= AUTO_TRIGGER_DEBOUNCE):
            self._last_sentence_end = 0.0
            self._generate()
        self.root.after(500, self._check_auto_trigger)

    # ── Generate ──────────────────────────────────────────────────────────────
    def _generate(self):
        global _current_profile
        p = _selected_provider.get()
        if p != "Ollama" and not self._api_key:
            self._status_var.set("請先連線（輸入 API Key）")
            return
        if self._generating:
            return
        ja_block, zh_block = _get_recent_context()
        if not ja_block:
            self._status_var.set("尚無語音資料，等對方說話後再按")
            return
        self._generating = True
        self._gen_btn.config(text="⏳ 生成中…", state="disabled", bg="#374151")
        self._set_text(self._reply_ja, "")
        self._set_text(self._reply_zh, "")
        system = _current_profile.get("system_prompt", DEFAULT_PROFILE["system_prompt"])
        self._status_var.set(
            f"呼叫 {p}/{_selected_model.get()}… 情境：{_current_profile['name']}")

        with _history_lock:
            hist_snapshot = list(_conversation_history)

        def run():
            try:
                ja_reply, zh_reply = call_ai(
                    self._api_key, p, _selected_model.get(),
                    ja_block, zh_block, system, hist_snapshot)
                _ui_queue.put({"type": "reply_done",
                               "ja_reply": ja_reply, "zh_reply": zh_reply,
                               "ja_block": ja_block, "zh_block": zh_block})
                _ui_queue.put({"type": "status", "text": "✓ 完成"})
            except Exception as e:
                _ui_queue.put({"type": "status", "text": f"錯誤：{e}"})
                _ui_queue.put({"type": "reply_done",
                               "ja_reply": f"[錯誤] {e}", "zh_reply": "",
                               "ja_block": ja_block, "zh_block": zh_block})
            finally:
                _ui_queue.put({"type": "gen_done"})

        threading.Thread(target=run, daemon=True).start()

    # ── TTS ───────────────────────────────────────────────────────────────────
    def _play_reply_tts(self, lang: str):
        text = self._last_ja_reply if lang == "ja" else self._last_zh_reply
        play_tts(text, lang)

    # ── Transcript helpers ────────────────────────────────────────────────────
    def _transcript_add(self, ts: str, ja: str, zh: str):
        line_no = self._trans_line
        self._ts_linemap[ts] = line_no
        self._trans_line += 1
        for widget, text, tag in (
            (self._trans_ja, ja,        "text_ja"),
            (self._trans_zh, zh or "…", "text_zh"),
        ):
            widget.config(state="normal")
            widget.insert("end", f"[{ts}]\n", "ts")
            widget.insert("end", text + "\n\n", tag)
            widget.see("end")
            widget.config(state="disabled")

    def _transcript_update_zh(self, ts: str, zh: str):
        if ts not in self._ts_linemap:
            return
        line_no = self._ts_linemap[ts]
        target  = f"{line_no * 3 - 1}.0"
        end_pos = f"{line_no * 3}.0"
        w = self._trans_zh
        w.config(state="normal")
        if w.get(target, end_pos).strip() == "…":
            w.delete(target, end_pos)
            w.insert(target, zh + "\n", "text_zh")
        w.config(state="disabled")

    # ── Export ────────────────────────────────────────────────────────────────
    def _export_transcript(self):
        if self._reply_sessions:
            records = self._reply_sessions
        else:
            with _lines_lock:
                records = [{"ts": ts, "ja": ja, "zh": zh,
                            "ja_reply": "", "zh_reply": ""}
                           for ts, ja, zh in _recent_lines]
        if not records:
            self._status_var.set("尚無資料可匯出")
            return

        filetypes = []
        default_ext = ".txt"
        if HAS_OPENPYXL:
            filetypes.append(("Excel 檔案", "*.xlsx"))
            default_ext = ".xlsx"
        if HAS_DOCX:
            filetypes.append(("Word 檔案", "*.docx"))
        filetypes.append(("文字檔案", "*.txt"))

        path = filedialog.asksaveasfilename(
            title="匯出逐字稿",
            defaultextension=default_ext,
            filetypes=filetypes,
            initialfile=f"翻譯記錄_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            parent=self.root)
        if not path:
            return
        try:
            if path.endswith(".xlsx"):
                export_excel(path, records)
            elif path.endswith(".docx"):
                export_word(path, records)
            else:
                lines = []
                for r in records:
                    lines.append(f"[{r.get('ts','')}]")
                    if r.get("ja"):       lines.append(f"日文：{r['ja']}")
                    if r.get("zh"):       lines.append(f"中文：{r['zh']}")
                    if r.get("ja_reply"): lines.append(f"AI回覆：{r['ja_reply']}")
                    if r.get("zh_reply"): lines.append(f"草稿：{r['zh_reply']}")
                    lines.append("")
                Path(path).write_text("\n".join(lines), encoding="utf-8")
            self._status_var.set(f"✓ 已匯出：{os.path.basename(path)}")
        except Exception as e:
            self._status_var.set(f"匯出失敗：{e}")

    # ── Poll ──────────────────────────────────────────────────────────────────
    def _poll(self):
        try:
            while True:
                msg = _ui_queue.get_nowait()
                t   = msg.get("type")
                if t == "status":
                    self._status_var.set(msg["text"])
                elif t == "transcript_add":
                    self._transcript_add(msg["ts"], msg["ja"], msg["zh"])
                    if msg.get("has_sentence_end") and self._auto_reply_var.get():
                        self._last_sentence_end = time.monotonic()
                elif t == "transcript_update_zh":
                    self._transcript_update_zh(msg["ts"], msg["zh"])
                elif t == "reply_chunk":
                    # Show raw streaming text as ja-box preview during generation
                    self._set_text(self._reply_ja, msg["text"])
                elif t == "reply_done":
                    ja_reply = msg["ja_reply"]
                    zh_reply = msg["zh_reply"]
                    self._last_ja_reply = ja_reply
                    self._last_zh_reply = zh_reply
                    self._set_text(self._reply_ja, ja_reply)
                    self._set_text(self._reply_zh, zh_reply)
                    # Update conversation history (keep last MAX_HISTORY_TURNS turns)
                    user_msg = _build_user_message(msg["ja_block"], msg["zh_block"])
                    asst_msg = json.dumps(
                        {"ja_reply": ja_reply, "zh_reply": zh_reply},
                        ensure_ascii=False)
                    with _history_lock:
                        _conversation_history.append({"role": "user",      "content": user_msg})
                        _conversation_history.append({"role": "assistant",  "content": asst_msg})
                        if len(_conversation_history) > MAX_HISTORY_TURNS * 2:
                            del _conversation_history[:-MAX_HISTORY_TURNS * 2]
                    self._history_count_var.set(f"歷史：{len(_conversation_history)//2} 輪")
                    # Store for export
                    with _lines_lock:
                        cutoff = (datetime.now() - timedelta(minutes=MAX_MINUTES)).strftime("%H:%M:%S")
                        recent = [(ts, ja, zh) for ts, ja, zh in _recent_lines
                                  if ts >= cutoff and ja][-MAX_SENTENCES:]
                    for ts, ja, zh in recent:
                        self._reply_sessions.append({
                            "ts": ts, "ja": ja, "zh": zh,
                            "ja_reply": ja_reply, "zh_reply": zh_reply
                        })
                    # Auto TTS
                    if self._tts_auto_var.get():
                        lang     = self._tts_lang_var.get()
                        tts_text = ja_reply if lang == "ja" else zh_reply
                        play_tts(tts_text, lang)
                elif t == "gen_done":
                    self._generating = False
                    self._gen_btn.config(
                        text=f"▶  生成回覆建議（最近 {MAX_SENTENCES} 句 / {MAX_MINUTES} 分鐘）",
                        state="normal", bg="#059669")
                elif t == "update_available":
                    ver, url, notes = msg["version"], msg["url"], msg["notes"]
                    body = (f"偵測到新版本 v{ver}，是否立即下載並更新？\n"
                            f"程式更新後會自動重啟。（jtw 模型不受影響）")
                    if notes:
                        body += f"\n\n更新說明：{notes}"
                    if messagebox.askyesno(f"新版本 v{ver} 可用", body, parent=self.root):
                        self._apply_update(url)
                elif t == "_do_restart":
                    subprocess.Popen(
                        ["cmd.exe", "/c", msg["bat"]],
                        creationflags=(subprocess.DETACHED_PROCESS
                                       | subprocess.CREATE_NEW_PROCESS_GROUP),
                        close_fds=True,
                    )
                    self.root.destroy()
                    sys.exit(0)
        except queue.Empty:
            pass
        self.root.after(150, self._poll)

    # ── Helpers ───────────────────────────────────────────────────────────────
    def _set_text(self, widget, text):
        widget.config(state="normal")
        widget.delete("1.0", "end")
        widget.insert("1.0", text)
        widget.config(state="disabled")

    def _copy_reply(self, lang: str = "ja"):
        text  = self._last_ja_reply if lang == "ja" else self._last_zh_reply
        label = "日文" if lang == "ja" else "中文"
        if not text:
            return
        if pyperclip:
            pyperclip.copy(text)
        else:
            self.root.clipboard_clear()
            self.root.clipboard_append(text)
        self._status_var.set(f"已複製{label}至剪貼簿 ✓")

    def _clear(self):
        self._set_text(self._reply_ja, "")
        self._set_text(self._reply_zh, "")
        self._last_ja_reply = ""
        self._last_zh_reply = ""
        for w in (self._trans_ja, self._trans_zh):
            self._set_text(w, "")
        self._ts_linemap.clear()
        self._trans_line = 1
        self._reply_sessions.clear()
        self._last_sentence_end = 0.0

    def _apply_update(self, url: str):
        """Download SEI update zip, write bat that replaces files and restarts."""
        import zipfile

        self._status_var.set("準備下載更新…")
        self._gen_btn.config(state="disabled")

        def _worker():
            try:
                tmp_dir = Path(tempfile.mkdtemp(prefix="sei_update_"))
                zip_path = tmp_dir / "sei_update.zip"
                extract_dir = tmp_dir / "extracted"

                def _progress(count, block_size, total):
                    if total > 0:
                        pct = min(100, count * block_size * 100 // total)
                        _ui_queue.put({"type": "status", "text": f"下載更新中… {pct}%"})

                urllib.request.urlretrieve(url, zip_path, reporthook=_progress)

                _ui_queue.put({"type": "status", "text": "解壓縮中…"})
                with zipfile.ZipFile(zip_path, "r") as z:
                    z.extractall(extract_dir)

                # If zip wraps a single folder, use that as root
                items = list(extract_dir.iterdir())
                src_dir = items[0] if len(items) == 1 and items[0].is_dir() else extract_dir

                bat_path = APP_DIR / "_sei_update.bat"
                exe_path = APP_DIR / "SEI_Reply_Assistant.exe"
                bat_lines = [
                    "@echo off",
                    "timeout /t 3 /nobreak >nul",
                    f'xcopy /e /y /q "{src_dir}\\*" "{APP_DIR}\\"',
                    f'start "" "{exe_path}"',
                    'del "%~f0"',
                ]
                bat_path.write_bytes(
                    "\r\n".join(bat_lines).encode("gbk", errors="replace")
                )

                _ui_queue.put({"type": "status", "text": "安裝中，程式即將重啟…"})
                _ui_queue.put({"type": "_do_restart", "bat": str(bat_path)})

            except Exception as e:
                _ui_queue.put({"type": "status", "text": f"更新失敗：{e}"})
                _ui_queue.put({"type": "gen_done"})

        threading.Thread(target=_worker, daemon=True).start()

    def _clear_history(self):
        with _history_lock:
            _conversation_history.clear()
        self._history_count_var.set("歷史：0 輪")
        self._status_var.set("對話歷史已清除")


def main():
    root = tk.Tk()
    App(root)
    root.mainloop()


if __name__ == "__main__":
    main()
